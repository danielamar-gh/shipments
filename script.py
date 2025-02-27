import pandas as pd
from docx import Document
from datetime import datetime
import os
from docx2pdf import convert
import requests
import urllib.parse
import time  # הוספת ספריית time להשהייות

# 1. קריאת הנתונים מקובץ CSV עם תמיכה בקידודים שונים ודילוג על שורות מיותרות
def load_data(csv_file):
    encodings = ["utf-8", "ISO-8859-8", "latin1", "windows-1255"]
    for enc in encodings:
        try:
            df = pd.read_csv(csv_file, encoding=enc, skiprows=2)
            df.columns = df.columns.str.strip()
            print(f"Successfully loaded file with encoding: {enc}")
            print("Columns in CSV:", df.columns.tolist())
            return df
        except UnicodeDecodeError:
            continue
    raise ValueError("Failed to decode CSV file with supported encodings")

# 2. סיכום נתוני המשלוחים בניכוי החזרות
def summarize_shipments(df):
    column_mapping = {
        'סוג מסמך': 'Document Type',
        'שם חשבון במסמך': 'Customer Name',
        'שם פריט': 'Item Name',
        'תאריך': 'Date',
        'כמות': 'Quantity',
        'מספר טלפון': 'Phone Number',
        'אסמכתא': 'Reference'
    }
    
    df.rename(columns=column_mapping, inplace=True)
    required_columns = list(column_mapping.values())
    for col in required_columns:
        if col not in df.columns:
            raise ValueError(f"Missing required column: {col}. Found columns: {df.columns.tolist()}")
    
    # המרת מספרי טלפון למחרוזת ועיצוב מתאים
    df['Phone Number'] = df['Phone Number'].astype(str).apply(lambda x: x.split('.')[0])  # הסרת הנקודה העשרונית
    
    # המרת אסמכתא למספר שלם ללא נקודה עשרונית
    df['Reference'] = df['Reference'].astype(str).apply(lambda x: x.split('.')[0])
    
    # המרת כמות למספר חיובי
    df['Quantity'] = df['Quantity'].abs()
    
    df['Date'] = pd.to_datetime(df['Date'], format='%d/%m/%Y', errors='coerce').dt.strftime('%d/%m/%Y')
    today_date = datetime.today().strftime('%d/%m/%Y')
    df = df[(df['Date'] == today_date) & ((df['Document Type'] == 'תעודת משלוח') | (df['Document Type'] == 'החזרה'))]
    
    summary = df.groupby(['Date', 'Document Type', 'Customer Name', 'Item Name', 'Phone Number', 'Quantity', 'Reference']).size().reset_index(name='count')
    summary = summary.drop('count', axis=1)
    summary.sort_values(by=['Document Type'], ascending=False, inplace=True)
    
    print("Summary Data:")
    print(summary)
    
    if summary.empty:
        print("⚠️ Warning: The summary DataFrame is empty. No data to insert into the table.")
    
    return summary

# 3. פונקציה להחלפת טקסט תוך שמירה על העיצוב
def replace_text_across_runs(paragraph, placeholder, replacement_text):
    full_text = "".join(run.text for run in paragraph.runs)
    if placeholder in full_text:
        full_text = full_text.replace(placeholder, replacement_text)
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = full_text  # כתיבה מחדש תוך שמירה על העיצוב

def send_text_message(chat_id, message):
    url = "https://7105.api.greenapi.com/waInstance7105171890/sendMessage/c82775d200444e42a2aadffb07574450300124341df4429ead"
    
    payload = {
        "chatId": chat_id,
        "message": message
    }
    
    headers = {
        'Content-Type': 'application/json'
    }
    
    try:
        response = requests.post(url, json=payload, headers=headers)
        print(f"Message send response: {response.text}")
        return True
    except Exception as e:
        print(f"Error sending message: {str(e)}")
        return False

def send_file_via_whatsapp(file_path, chat_id):
    url = "https://7105.api.greenapi.com/waInstance7105171890/sendFileByUpload/c82775d200444e42a2aadffb07574450300124341df4429ead"
    
    # שימוש בשם הקובץ המקורי ללא הסיומת .pdf
    original_filename = os.path.basename(file_path)
    filename_without_ext = os.path.splitext(original_filename)[0]
    
    payload = {
        'chatId': chat_id,
        'fileName': filename_without_ext,  # שימוש בשם הקובץ ללא סיומת
        'caption': 'שלום מצורף הדוח היומי'  # הוספת הודעת הטקסט כcaption
    }
    
    with open(file_path, 'rb') as file:
        files = [
            ('file', (original_filename, file, 'application/pdf'))
        ]
        
        headers = {
            'Accept': '*/*',
            'Accept-Encoding': 'gzip, deflate, br',
        }
        
        response = requests.post(url, data=payload, files=files, headers=headers)
        print(f"WhatsApp API Response for {filename_without_ext}:", response.text)
        return response.status_code == 200

def send_whatsapp_content(chat_id, customer, pdf_file):
    # שליחת הודעת פתיחה
    greeting_message = "שלום מצורף הדוח היומי"
    message_sent = False
    retries = 3  # מספר ניסיונות לשליחת ההודעה
    
    # ניסיון לשלוח את ההודעה מספר פעמים
    for attempt in range(retries):
        if send_text_message(chat_id, greeting_message):
            message_sent = True
            print(f"הודעת פתיחה נשלחה בהצלחה ל-{customer}")
            break
        else:
            print(f"ניסיון {attempt + 1} נכשל, מנסה שוב...")
            time.sleep(2)  # המתנה של 2 שניות בין ניסיונות
    
    if not message_sent:
        print(f"נכשל לשלוח הודעת פתיחה ל-{customer} אחרי {retries} ניסיונות")
        return False
    
    # המתנה קצרה בין שליחת ההודעה לקובץ
    time.sleep(3)
    
    # שליחת הקובץ
    if send_file_via_whatsapp(pdf_file, chat_id):
        print(f"הקובץ {pdf_file} נשלח בהצלחה ל-WhatsApp למספר {chat_id}")
        return True
    else:
        print(f"שגיאה בשליחת הקובץ {pdf_file} ל-WhatsApp למספר {chat_id}")
        return False

# 4. יצירת קובץ PDF לכל לקוח
def fill_word_template_per_customer(summary, template_path, output_dir):
    os.makedirs(output_dir, exist_ok=True)
    unique_customers = summary[['Customer Name', 'Phone Number']].drop_duplicates()
    current_date = datetime.today().strftime('%d_%m_%Y')
    
    for _, customer_info in unique_customers.iterrows():
        customer = customer_info['Customer Name']
        phone = customer_info['Phone Number']
        
        # בדיקה והכנת מספר הטלפון לפורמט המתאים
        phone_str = str(phone).strip()
        if phone_str.startswith('0'):  # אם המספר מתחיל ב-0
            phone_str = '972' + phone_str[1:]  # החלפת ה-0 ב-972
        elif not phone_str.startswith('972'):  # אם המספר לא מתחיל ב-972
            phone_str = '972' + phone_str
            
        print(f"Preparing to send to phone number: {phone_str}")
        whatsapp_chat_id = f'{phone_str}@c.us'
        
        customer_data = summary[summary['Customer Name'] == customer]
        doc = Document(template_path)
        
        # מילוי שדות דינמיים תוך שמירה על העיצוב
        for paragraph in doc.paragraphs:
            replace_text_across_runs(paragraph, "{שם חשבון במסמך}", customer)
            replace_text_across_runs(paragraph, "{תאריך}", datetime.today().strftime('%d-%m-%Y'))
        
        table = doc.tables[0]
        headers = [cell.text.strip().replace("\u200e", "").replace("\ufeff", "").replace("{", "").replace("}", "") for cell in table.rows[0].cells]
        print(f"Processing customer: {customer} (Phone: {phone_str})")
        print("Table Headers:", headers)
        
        column_map = {
            "תאריך": "Date",
            "סוג מסמך": "Document Type",
            "שם חשבון במסמך": "Customer Name",
            "שם פריט": "Item Name",
            "כמות": "Quantity",
            "אסמכתא": "Reference"
        }
        
        column_indices = {column_map[key]: headers.index(key) for key in headers if key in column_map}
        
        for _, row in customer_data.iterrows():
            new_row = table.add_row().cells
            for col_name, df_col in column_indices.items():
                new_row[df_col].text = str(row[col_name])
        
        # שמירת הקבצים עם שמות בעברית
        filename_base = f"דוח_{current_date}_{customer}"
        word_file = os.path.join(output_dir, f"{filename_base}.docx")
        pdf_file = os.path.join(output_dir, f"{filename_base}.pdf")
        
        doc.save(word_file)
        convert(word_file)
        os.remove(word_file)  # מחיקת קובץ ה-Word לאחר יצירת ה-PDF
        print(f"מסמך PDF נשמר: {pdf_file}")
        
        # שליחת הקובץ באמצעות WhatsApp עם הודעת טקסט בתור caption
        if send_file_via_whatsapp(pdf_file, whatsapp_chat_id):
            print(f"הקובץ {pdf_file} נשלח בהצלחה ל-WhatsApp למספר {phone_str}")
        else:
            print(f"שגיאה בשליחת הקובץ {pdf_file} ל-WhatsApp למספר {phone_str}")
        
        # המתנה קצרה בין לקוחות
        time.sleep(2)

# דוגמה לשימוש
csv_file = 'shipments.csv'
template_path = 'Template.docx'
output_dir = os.path.join(os.path.dirname(__file__), 'customer_reports')

df = load_data(csv_file)
summary = summarize_shipments(df)
fill_word_template_per_customer(summary, template_path, output_dir)


